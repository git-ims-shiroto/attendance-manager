from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_FILE = OUT_DIR / "勤怠管理プラグイン_利用者向け機能ガイド_納品版.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5F6B76"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "2C5F2E"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B7C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_font(run, size=None, bold=None, color=None, italic=None, east_asia="Yu Gothic"):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=None, bold=None, color=None):
    for run in paragraph.runs:
        set_font(run, size=size, bold=bold, color=color)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, num_fmt, text, left=540, hanging=270):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "Yu Gothic")
    r_pr.append(r_fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_body(doc, text, bold_prefix=None, color=None, after=6):
    p = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=color or INK)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, color=color)
    else:
        r = p.add_run(text)
        set_font(r, color=color)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_bullet(doc, text, bullet_num_id):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, bullet_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p


def add_step(doc, title, description, decimal_num_id):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, decimal_num_id)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run(title + "　")
    set_font(r1, bold=True, color=INK)
    r2 = p.add_run(description)
    set_font(r2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_paragraph_font(p)
    set_keep_with_next(p)
    return p


def add_callout(doc, label, text, color=INK, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_row_cant_split(table.rows[0])
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D5DEE7", size="5")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label + "　")
    set_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_font(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers, rows, widths, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_row_cant_split(hdr)
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=10, bold=True, color=INK)
    for row_data in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, text in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = (aligns[i] if aligns else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(text))
            set_font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section):
    def fill_header(header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("勤怠管理プラグイン  |  利用者向け機能ガイド")
        set_font(r, size=8.5, color=MUTED)

    def fill_footer(footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run("- ")
        set_font(r, size=9, color=MUTED)
        add_field(p, "PAGE")
        r = p.add_run(" -")
        set_font(r, size=9, color=MUTED)

    for header in (section.header, section.even_page_header, section.first_page_header):
        fill_header(header)
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        fill_footer(footer)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    configure_styles(doc)
    for section in doc.sections:
        configure_page(section)
    bullet_num_id = add_numbering_definition(doc, "bullet", "●")

    section = doc.sections[0]

    # Cover: editorial_cover pattern, adapted for a practical Japanese operator guide.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("利用者向けガイド")
    set_font(r, size=12, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("勤怠管理プラグイン")
    set_font(r, size=30, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("画面でできることと、基本的な使い方")
    set_font(r, size=15, color=DARK_BLUE)

    add_callout(doc, "このガイドについて", "勤怠管理の画面を使う方に向けて、機能と操作だけをやさしい言葉でまとめています。", fill="FFF8E8", color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("有限会社たんぽぽ運送")
    set_font(r, size=12, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026年8月版")
    set_font(r, size=10, color=MUTED)

    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    configure_page(body_section)
    body_section.header.is_linked_to_previous = False
    body_section.footer.is_linked_to_previous = False
    body_section.even_page_header.is_linked_to_previous = False
    body_section.first_page_header.is_linked_to_previous = False
    body_section.even_page_footer.is_linked_to_previous = False
    body_section.first_page_footer.is_linked_to_previous = False
    body_section.different_first_page_header_footer = False
    configure_header_footer(body_section)

    add_heading(doc, "1. このプラグインでできること", 1)
    add_body(doc, "社員ごとの1か月の勤怠を確認し、必要な箇所を直して保存できます。また、全社員の集計確認や、休日・職種の設定も行えます。")
    add_table(doc,
              ["メニュー", "できること"],
              [
                  ("長距離", "長距離ドライバーの日ごとの勤務内容、週ごとの合計、月の合計を確認・更新します。"),
                  ("地場・事務", "地場勤務・事務の社員の日ごとの勤務内容、週ごとの合計、月の合計を確認・更新します。"),
                  ("集計一覧", "選んだ月について、全社員の出勤日数や労働時間などを一覧で確認します。"),
                  ("設定", "所属ごとの休日と、職種をどちらの管理表に表示するかを設定します。"),
              ],
              [1800, 7560])
    add_callout(doc, "表示について", "利用できるメニューは、利用者に与えられている権限によって異なります。「設定」は設定を担当する方だけに表示されます。")

    add_heading(doc, "2. 基本の使い方", 1)
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.")
    add_step(doc, "メニューを開く", "WordPressの管理画面で「勤怠管理」を選び、目的の画面を開きます。", decimal_num_id)
    add_step(doc, "社員と月を選ぶ", "必要に応じて所属で絞り込み、社員名と対象月を選んで「表示」を押します。", decimal_num_id)
    add_step(doc, "内容を確認する", "日ごとの内容、週ごとの合計、月の合計を順に確認します。画面に注意の表示がある場合は、内容も確認します。", decimal_num_id)
    add_step(doc, "必要な箇所を直す", "勤怠種別、地場・長距離の印、早退・遅刻時間、備考を必要に応じて変更します。", decimal_num_id)
    add_step(doc, "保存する", "「保存（更新）」を押し、「保存しました」と表示されたことを確認します。", decimal_num_id)
    add_callout(doc, "大切", "画面上で変更しただけでは保存されません。作業の最後に必ず「保存（更新）」を押してください。", color=RED, fill="FFF2F2")

    add_heading(doc, "3. 長距離の勤怠を確認する", 1)
    add_body(doc, "「長距離」画面では、長距離ドライバーの1か月分の勤怠を確認できます。")
    add_heading(doc, "社員と月を表示する", 2)
    add_bullet(doc, "所属のボタンが表示されている場合は、先に所属を選ぶと社員名を絞り込めます。", bullet_num_id)
    add_bullet(doc, "「社員名」と「対象月」を選び、「表示」を押します。", bullet_num_id)
    add_bullet(doc, "表示後は、表の見出しに社員名と対象月が表示されます。左右の矢印で前月・翌月へ移動できます。", bullet_num_id)
    add_heading(doc, "日ごとの表で確認できる内容", 2)
    add_table(doc,
              ["項目", "意味"],
              [
                  ("勤怠種別", "出勤、休日、有給、欠勤など、その日の扱いです。"),
                  ("始業時刻・終業時刻", "仕事を始めた時刻と終えた時刻です。"),
                  ("拘束時間", "始業から終業までの全体時間です。"),
                  ("労働時間", "休憩などを除いた、働いた時間です。"),
                  ("運転時間・積卸時間", "運転した時間と、荷物の積み下ろしをした時間です。"),
                  ("休憩時間・深夜時間", "休憩した時間と、深夜に働いた時間です。"),
                  ("日残業", "その日について計算された残業時間です。"),
                  ("地場", "その日の勤務を地場勤務として扱うときにオンにします。"),
                  ("早退／遅刻", "早退または遅刻した時間を「0:30」のように入力します。"),
                  ("備考", "補足しておきたい内容を入力します。"),
              ],
              [2400, 6960])
    add_callout(doc, "休日出勤の表示", "休日に勤務がある日は、「法定休出勤」または「所定休出勤」の印が日付の近くに表示されます。")

    add_heading(doc, "4. 地場・事務の勤怠を確認する", 1)
    add_body(doc, "「地場・事務」画面では、地場勤務・事務の社員の1か月分の勤怠を確認できます。社員と月の選び方、表の見方、保存方法は「長距離」画面と同じです。")
    add_heading(doc, "長距離画面との違い", 2)
    add_body(doc, "日ごとの表に「長距離」という切り替えがあります。その日の勤務を長距離勤務として扱うときにオンにします。")
    add_bullet(doc, "オン・オフを変更して保存すると、最新の内容で表が表示し直されます。", bullet_num_id)
    add_bullet(doc, "勤怠種別、早退／遅刻、備考も変更できます。", bullet_num_id)
    add_bullet(doc, "始業時刻や労働時間などは、記録されている内容をもとに画面へ表示されます。", bullet_num_id)
    add_callout(doc, "使い分け", "長距離の社員がその日だけ地場勤務をした場合は「長距離」画面の「地場」を使います。地場・事務の社員がその日だけ長距離勤務をした場合は「地場・事務」画面の「長距離」を使います。")

    add_heading(doc, "5. 日ごとの内容を直して保存する", 1)
    add_heading(doc, "変更できる項目", 2)
    add_bullet(doc, "勤怠種別：出勤、法定休、法定振替休、所定休、所定振替休、有給、欠勤、緊急出動、または「―」から選びます。", bullet_num_id)
    add_bullet(doc, "地場／長距離：その日の勤務内容に合わせてオン・オフを切り替えます。", bullet_num_id)
    add_bullet(doc, "早退／遅刻：時間を「時:分」で入力します。例：30分なら「0:30」。", bullet_num_id)
    add_bullet(doc, "備考：確認した人が分かるように、必要な補足を短く入力します。", bullet_num_id)
    add_heading(doc, "保存後の動き", 2)
    add_body(doc, "「保存（更新）」を押すと、保存件数が表示され、日ごとの表・週の合計・月の合計が新しい内容で更新されます。地場／長距離の切り替えを変更した場合は、保存後に画面全体が表示し直されます。")
    add_callout(doc, "入力の目安", "早退／遅刻は「0:30」「1:15」のように入力します。分だけの数字や文章は入れず、画面の「0:00」と同じ形にそろえてください。", color=GOLD, fill="FFF8E8")

    add_heading(doc, "6. 週の合計と月の合計を見る", 1)
    add_heading(doc, "週の合計", 2)
    add_body(doc, "日ごとの表の下に、期間ごとの合計が表示されます。期間、開始日、終了日、日数のほか、拘束・労働・運転・積卸・休憩・深夜・日残業・週残業・確定残業を確認できます。")
    add_bullet(doc, "「日残業」は、各日の残業を合計した時間です。", bullet_num_id)
    add_bullet(doc, "「週残業」は、1週間の合計から計算された残業時間です。", bullet_num_id)
    add_bullet(doc, "「確定残業」は、日残業と週残業をもとに最終的に確認する残業時間です。", bullet_num_id)
    add_bullet(doc, "月をまたぐ週は、前月からの分や翌月へ引き継ぐ分が分かるように表示されます。", bullet_num_id)
    add_heading(doc, "月の合計", 2)
    add_body(doc, "月の合計では、出勤日数、欠勤日数、休日出勤日数、有給消化日数、有給残日数、労働時間、早退遅刻時間、確定残業時間を確認できます。")
    add_callout(doc, "確認の順番", "まず日ごとの勤怠種別と時間を確認し、次に週の合計、最後に月の合計を見ると確認しやすくなります。")

    add_heading(doc, "7. 全社員の集計一覧を見る", 1)
    add_body(doc, "「集計一覧」では、選んだ月の全社員分を1つの表で確認できます。長距離と地場・事務の両方が対象です。")
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.")
    add_step(doc, "対象月を選ぶ", "確認したい月を選びます。", decimal_num_id)
    add_step(doc, "読み込む", "「読み込む」を押します。", decimal_num_id)
    add_step(doc, "一覧を確認する", "社員コード、社員名、出勤・欠勤・休日出勤・有給の日数、労働時間、早退遅刻時間、確定残業時間を確認します。", decimal_num_id)
    add_callout(doc, "この画面の役割", "全体を見渡すための画面です。内容を直すときは、その社員の「長距離」または「地場・事務」画面を開いて変更してください。")

    add_heading(doc, "8. 休日を設定する", 1)
    add_body(doc, "「設定」の「休日マスタ設定」では、所属ごとに、何曜日の第何週を所定休日にするかを登録します。この設定は長距離と地場・事務の両方で使われます。")
    add_heading(doc, "休日ルールを追加する", 2)
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.")
    add_step(doc, "所属を選ぶ", "休日を設定する所属を選びます。", decimal_num_id)
    add_step(doc, "曜日を選ぶ", "所定休日にする曜日を選びます。", decimal_num_id)
    add_step(doc, "対象週を入力する", "第2週と第4週なら「2,4」のように、半角のカンマで区切って入力します。", decimal_num_id)
    add_step(doc, "保存する", "「保存」を押し、登録済みルール一覧に表示されたことを確認します。", decimal_num_id)
    add_heading(doc, "登録済みルールを管理する", 2)
    add_bullet(doc, "編集：登録内容を入力欄に呼び出し、直して保存します。途中でやめる場合は「キャンセル」を押します。", bullet_num_id)
    add_bullet(doc, "無効化／有効化：削除せずに、一時的に使わない状態と、再び使う状態を切り替えます。", bullet_num_id)
    add_bullet(doc, "削除：不要になったルールを消します。確認画面で同意すると削除されます。", bullet_num_id)
    add_callout(doc, "注意", "休日の設定を変えると、その所属の勤怠表示に影響します。変更前に、対象の所属・曜日・週が正しいか確認してください。", color=RED, fill="FFF2F2")

    add_heading(doc, "9. 職種の表示先を設定する", 1)
    add_body(doc, "「設定」の「種別管理」では、職種ごとに社員を「長距離」と「地場・事務」のどちらの管理表へ表示するかを決めます。")
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.")
    add_step(doc, "職種を確認する", "一覧に表示されている職種名を確認します。", decimal_num_id)
    add_step(doc, "表示先を選ぶ", "その職種の「長距離」または「地場・事務」を選びます。", decimal_num_id)
    add_step(doc, "保存結果を確認する", "選ぶとすぐに保存され、状態が「設定済」に変わります。", decimal_num_id)
    add_bullet(doc, "「未設定」の職種は、どちらの管理表に表示するかがまだ決まっていません。", bullet_num_id)
    add_bullet(doc, "職種そのものが一覧にない場合は、先に社員情報を管理する画面で職種を登録します。", bullet_num_id)
    add_callout(doc, "変更後の確認", "職種の表示先を変えたら、「長距離」または「地場・事務」画面で、対象社員が正しい側に表示されることを確認してください。")

    add_heading(doc, "10. 画面の表示や保存で困ったとき", 1)
    add_table(doc,
              ["状況", "確認すること"],
              [
                  ("社員名を選べない", "所属の絞り込みを「すべて」に戻すか、職種の表示先設定を確認します。"),
                  ("表示ボタンを押せない", "社員名が選ばれているか、対象月が入っているか確認します。"),
                  ("変更が残っていない", "「保存（更新）」を押し、保存完了の表示を確認します。"),
                  ("保存に失敗したと表示される", "画面を更新してもう一度試します。続く場合は管理担当者へ連絡します。"),
                  ("通信エラーと表示される", "インターネット接続を確認し、少し待ってからもう一度保存します。"),
                  ("注意の表示が出ている", "表示された文章を確認し、元の勤務記録に不足や食い違いがないか管理担当者へ確認します。"),
                  ("設定メニューが見えない", "設定を変更できる権限がない可能性があります。設定担当者へ依頼します。"),
              ],
              [2600, 6760])
    add_callout(doc, "連絡するとき", "社員名、対象月、発生した操作、画面に表示された文章を伝えると、確認がスムーズです。")

    add_heading(doc, "日常の確認チェック", 1)
    for item in [
        "対象の社員と月が合っている",
        "日ごとの勤怠種別が合っている",
        "地場／長距離の切り替えが合っている",
        "早退・遅刻時間と備考が必要に応じて入っている",
        "週の合計と月の合計に不自然な点がない",
        "最後に「保存（更新）」を押した",
        "保存完了の表示を確認した",
    ]:
        add_bullet(doc, item, bullet_num_id)

    # Apply section settings and core properties.
    for sec in doc.sections:
        configure_page(sec)
    props = doc.core_properties
    props.title = "勤怠管理プラグイン 利用者向け機能ガイド"
    props.subject = "勤怠管理プラグインの利用者向け解説書"
    props.author = "有限会社たんぽぽ運送"
    props.keywords = "勤怠管理, 利用者ガイド, 長距離, 地場, 事務"

    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build()
